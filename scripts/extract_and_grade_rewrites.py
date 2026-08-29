#!/usr/bin/env python3
"""
Workflow B1 v2.1 - Rewrite Completion Grading (format-aware)
CHANGELOG 2026-06-07 (v2.1): + rewrite_marker format - keys off "My Rewrite for
  slide N:" anchor and ignores pasted feedback prose. B1->B2 JSON contract unchanged.

Based on deployed extract_and_grade_rewrites.py. Adds format detection so the
reader isolates the STUDENT's writing from the pasted email direction:

  * answer_marker  : email puts a "Your answer:" slot under each slide; student
                     text = everything after "Your answer:" in that section.
  * legacy_inline  : older emails; student text = the paragraph(s) under the
                     "Slide N" direction line (drop the direction line).
  * table_template : the "YOUR IMPROVED ANSWER / (No answer submitted)" doc -
                     order is unreliable via flat text; FLAG, do not score.

Completion denominator fixed at 5 (policy). Flags any doc whose section count
!= 5. Challenge tracked separately. Writes nothing to the dashboard.
"""
import re, argparse, json, csv
from datetime import date
from pathlib import Path
from docx import Document

SLIDE_MARKER_RE = re.compile(r"(?:^|\n)\s*(?:\d+[\.\)]\s*)?(?:#+\s*)?(?:\*\*)?Slide\s+(\d+)\b", re.IGNORECASE)
CHALLENGE_MARKER_RE = re.compile(r"(?:^|\n)\s*(?:#+\s*)?(?:\*\*)?Challenge\b", re.IGNORECASE)
ANSWER_MARKER_RE = re.compile(r"your answer\s*:", re.IGNORECASE)
MY_REWRITE_RE = re.compile(r"(?:^|\n)\s*(?:#+\s*)?(?:\*\*)?my rewrite for slide\s+(\d+)\s*:", re.IGNORECASE)
CLOSING_RE = re.compile(r"(?i)when you'?re ready")
ADDRESSED_WORD_THRESHOLD = 8   # 2026-07-06: 15->8; even a short on-topic answer shows effort (per K.)
EXPECTED_SECTIONS = 5
TABLE_SIGNS = ("improved answer", "no answer submitted", "your first answer")

# --- canonical lesson-name guard (mirror of lesson_name_guard.py; KEEP IN SYNC) ---
# Halts on a non-canonical lesson name so a mismatch never silently writes a
# blank Schoology import cell downstream. Restores the guard dropped in the
# 2026-06-07 marker-parser fork.
_CANON = ["Natural Selection", "Evidence of Evolution", "Introduction to Evolution",
          "Classifying Organisms", "Time of Death", "What is Death?"]
_OOS = {"pig autopsy", "stages of decomposition"}
def _norm_lesson(s):
    return re.sub(r"\s+", " ", (s or "").strip()).rstrip(".?!").casefold()
_LU = {_norm_lesson(c): c for c in _CANON}
def resolve_lesson(raw):
    n = _norm_lesson(raw)
    if n in _OOS:
        return None, f"OUT OF SCOPE lesson '{raw}' - do not process."
    c = _LU.get(n)
    if c is None:
        return None, (f"UNKNOWN LESSON '{raw}' - not canonical; writing it would "
                      f"blank the Schoology import cell. Add to canonical list + "
                      f"Dashboard column first.")
    return c, ("" if c == raw else f"normalized '{raw}' -> '{c}'")

def wc(t): return len(t.split())

def read_doc(path):
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs]
    tcells = [c.text.strip() for tb in doc.tables for r in tb.rows for c in r.cells]
    return doc, paras, tcells

def detect_format(paras, tcells):
    low = "\n".join(paras + tcells).lower()
    if any(s in low for s in TABLE_SIGNS):
        return "table_template"
    if "my rewrite for slide" in low:
        return "rewrite_marker"
    if "your answer:" in low:
        return "answer_marker"
    return "legacy_inline"

def split_sections(text):
    cm = CHALLENGE_MARKER_RE.search(text)
    cstart = cm.start() if cm else len(text)
    region = text[:cstart]
    markers = list(SLIDE_MARKER_RE.finditer(region))
    secs = []
    for i, m in enumerate(markers):
        s = m.start(); e = markers[i+1].start() if i+1 < len(markers) else len(region)
        secs.append((int(m.group(1)), region[s:e].strip()))
    ch = ""
    if cm:
        rest = text[cstart:].split("\n", 1)
        ch = rest[1].strip() if len(rest) > 1 else ""
        ch = CLOSING_RE.split(ch)[0].strip()
    return secs, ch

def split_rewrite_marker(text):
    """rewrite_marker: student labels each answer 'My Rewrite for slide N:'.
    Keys ONLY off that anchor so pasted feedback prose (which mentions 'slide N')
    is ignored. Segment IS the student's text (everything after the anchor)."""
    cm = CHALLENGE_MARKER_RE.search(text)
    cstart = cm.start() if cm else len(text)
    region = text[:cstart]
    markers = list(MY_REWRITE_RE.finditer(region))
    secs = []
    for i, m in enumerate(markers):
        s = m.end(); e = markers[i+1].start() if i+1 < len(markers) else len(region)
        secs.append((int(m.group(1)), region[s:e].strip()))
    ch = ""
    if cm:
        rest = text[cstart:].split("\n", 1)
        ch = rest[1].strip() if len(rest) > 1 else ""
        ch = CLOSING_RE.split(ch)[0].strip()
    return secs, ch

def student_text(segment, fmt):
    if fmt == "answer_marker":
        m = ANSWER_MARKER_RE.search(segment)
        return segment[m.end():].strip() if m else ""
    parts = segment.split("\n", 1)          # legacy: drop the "Slide N - direction" line
    return parts[1].strip() if len(parts) > 1 else ""

def parse(paras, tcells, fmt, docx_bytes=None):
    flags = []
    if fmt == "table_template":
        # Route the "YOUR IMPROVED ANSWER" table family through the canonical
        # aspect_extractor (docx table path), then map into B1 rewrite dicts.
        # Student text = the rewrite cell, falling back to the first-answer cell.
        try:
            import aspect_extractor
        except Exception as e:
            return [], None, ["table template - aspect_extractor unavailable (%s); needs review" % e]
        if not docx_bytes:
            return [], None, ["table template - no docx bytes to parse; needs review"]
        secs = aspect_extractor.extract_sections(docx_bytes=docx_bytes).get("sections", [])
        if not secs:
            return [], None, ["table template - no sections extracted; needs review"]
        rewrites = []
        for s in secs:
            st = (s.get("rewrite") or s.get("first") or "").strip()
            try:
                num = int(s.get("slide"))
            except (TypeError, ValueError):
                continue
            rewrites.append({"slide_number": num, "rewrite_text": st, "word_count": wc(st),
                             "addressed": wc(st) >= ADDRESSED_WORD_THRESHOLD})
        challenge = None
        distinct = {r["slide_number"] for r in rewrites}
        if len(distinct) != EXPECTED_SECTIONS:
            flags.append(f"expected 5 slide sections, found {len(distinct)}")
        return rewrites, challenge, flags
    text = "\n".join(paras + tcells)
    if fmt == "rewrite_marker":
        secs, ch_text = split_rewrite_marker(text)
        if not secs:
            return [], None, ["no 'My Rewrite for slide N:' markers found - unparseable"]
        rewrites = [{"slide_number": num, "rewrite_text": st, "word_count": wc(st),
                     "addressed": wc(st) >= ADDRESSED_WORD_THRESHOLD} for num, st in secs]
    else:
        secs, ch_text = split_sections(text)
        if not secs:
            return [], None, ["no Slide markers found - unparseable"]
        rewrites = []
        for num, seg in secs:
            st = student_text(seg, fmt)
            rewrites.append({"slide_number": num, "rewrite_text": st, "word_count": wc(st),
                             "addressed": wc(st) >= ADDRESSED_WORD_THRESHOLD})
    challenge = None
    if ch_text:
        challenge = {"text": ch_text, "word_count": wc(ch_text), "addressed": wc(ch_text) >= ADDRESSED_WORD_THRESHOLD}
    distinct = {r["slide_number"] for r in rewrites}
    if len(distinct) != EXPECTED_SECTIONS:
        flags.append(f"expected 5 slide sections, found {len(distinct)}")
    return rewrites, challenge, flags

def score(rewrites):
    addressed = {r["slide_number"] for r in rewrites if r["addressed"]}
    n = min(len(addressed), EXPECTED_SECTIONS)
    return n, round(100.0 * n / EXPECTED_SECTIONS, 1)

def grade_file(path, fmt_override=None):
    doc, paras, tcells = read_doc(path)
    fmt = fmt_override or detect_format(paras, tcells)
    _docx_bytes = None
    if fmt == "table_template":
        with open(path, "rb") as _fh:
            _docx_bytes = _fh.read()
    rewrites, challenge, flags = parse(paras, tcells, fmt, _docx_bytes)
    n, pct = score(rewrites)
    return {"format": fmt, "rewrites": rewrites, "challenge": challenge,
            "addressed": n, "completion_percent": pct, "flags": flags}

# ---------------------------------------------------------------------------
# Output layer (preserves the v1 per-student JSON schema consumed by B2)
# ---------------------------------------------------------------------------
import sys

def _completion(g):
    return {
        "rewrite_sections_addressed": g["addressed"],
        "total_rewrite_sections_expected": EXPECTED_SECTIONS,
        "completion_percent": g["completion_percent"],
        "challenge_addressed": bool(g["challenge"] and g["challenge"]["addressed"]),
        "total_word_count": sum(r["word_count"] for r in g["rewrites"]) + (g["challenge"]["word_count"] if g["challenge"] else 0),
    }

def write_student_json(out_dir, student, lesson, klass, g):
    comp = _completion(g)
    payload = {
        "student": student, "lesson": lesson, "class": klass,
        "date": date.today().isoformat(),
        "format": g["format"], "flags": g["flags"],          # additive keys; B2 ignores unknown
        "completion": comp,
        "rewrites": g["rewrites"],
        "challenge_response": g["challenge"],
    }
    p = out_dir / "rewrites" / f"{student}.json"
    p.parent.mkdir(parents=True, exist_ok=True)
    with open(p, "w") as f:
        json.dump(payload, f, indent=2)
    return comp

def write_completion_csv(out_dir, rows):
    with open(out_dir / "rewrite_completion_report.csv", "w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=["student","format","rewrite_sections_addressed",
            "total_rewrite_sections_expected","completion_percent","challenge_addressed",
            "total_word_count","flags"])
        w.writeheader(); w.writerows(rows)

def write_summary_md(out_dir, rows, lesson, klass):
    L = ["# Workflow B1 v2 â Rewrite Completion Report", "",
         f"**Lesson:** `{lesson}`", f"**Class:** `{klass}`",
         f"**Date:** {date.today().isoformat()}", f"**Rewrites processed:** {len(rows)}", "",
         "| Student | Format | Addressed | % | Challenge | Flags |",
         "|---|---|---|---|---|---|"]
    for r in sorted(rows, key=lambda x: x["student"].lower()):
        ch = "yes" if r["challenge_addressed"] else "no"
        L.append(f"| {r['student']} | {r['format']} | {r['rewrite_sections_addressed']}/{r['total_rewrite_sections_expected']} | {r['completion_percent']}% | {ch} | {r['flags']} |")
    L += ["", "Per-student JSON in `rewrites/` is the Workflow B2 input (schema unchanged)."]
    with open(out_dir / "workflow_b1_summary.md", "w") as f:
        f.write("\n".join(L))

def run(draft_dir, output_dir, lesson, klass):
    draft_dir = Path(draft_dir); output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    canon, _note = resolve_lesson(lesson)
    if canon is None:
        print("HALT: " + _note, file=sys.stderr); sys.exit(3)
    if _note:
        print(_note, file=sys.stderr)
    lesson = canon
    files = [p for p in sorted(draft_dir.glob("*.docx")) if re.match(r"(?i)^(rewrite|draft)", p.stem)]
    if not files:
        print(f"No REWRITE/Draft .docx found in {draft_dir}", file=sys.stderr); sys.exit(2)
    pref = re.compile(rf"(?i)^(rewrite\.notes\.{re.escape(lesson)}|draft)[\s._-]*")
    rows = []
    for df in files:
        student = pref.sub("", df.stem).strip() or df.stem
        g = grade_file(str(df))
        comp = write_student_json(output_dir, student, lesson, klass, g)
        rows.append({"student": student, "format": g["format"],
            "rewrite_sections_addressed": comp["rewrite_sections_addressed"],
            "total_rewrite_sections_expected": comp["total_rewrite_sections_expected"],
            "completion_percent": comp["completion_percent"],
            "challenge_addressed": comp["challenge_addressed"],
            "total_word_count": comp["total_word_count"],
            "flags": "; ".join(g["flags"])})
    write_completion_csv(output_dir, rows); write_summary_md(output_dir, rows, lesson, klass)
    print(f"Processed {len(rows)} rewrite docs -> {output_dir}")
    flagged = [r for r in rows if r["flags"]]
    if flagged: print(f"{len(flagged)} doc(s) flagged for review")

def main():
    ap = argparse.ArgumentParser(description="Workflow B1 v2 - rewrite completion grading (format-aware)")
    ap.add_argument("draft_dir"); ap.add_argument("output_dir")
    ap.add_argument("--lesson", required=True)
    ap.add_argument("--class", dest="klass", required=True)
    a = ap.parse_args(); run(a.draft_dir, a.output_dir, a.lesson, a.klass)

if __name__ == "__main__":
    main()
